from docx import Document

# Crear un documento de Word
doc = Document()

# Agregar título
doc.add_heading('Curso MOOC: Ética, Bioética de la Investigación e Integridad Científica', level=1)

# Introducción
doc.add_heading('Introducción al Curso', level=2)
doc.add_paragraph(
    "Este curso está diseñado para proporcionar a los estudiantes una comprensión profunda de los principios éticos y bioéticos que guían la investigación científica, así como la importancia de la integridad en el proceso investigativo. Los participantes explorarán temas relacionados con la ética en la investigación con seres vivos y no vivos, las garantías para la participación de seres humanos en investigaciones, y la protección de datos y confidencialidad."
)

# Objetivos del curso
doc.add_heading('Objetivos del Curso', level=2)
doc.add_paragraph(
    "Al finalizar este curso, los estudiantes serán capaces de:"
)
objectives = [
    "Comprender los principios fundamentales de la ética y la bioética en la investigación.",
    "Aplicar estos principios en diversos contextos de investigación científica.",
    "Evaluar críticamente casos específicos de investigación para identificar y resolver tensiones éticas y bioéticas.",
    "Utilizar el método deliberativo para alcanzar consensos en situaciones de conflicto ético.",
    "Reconocer la importancia de la integridad científica y aplicar buenas prácticas en sus propias investigaciones."
]
for obj in objectives:
    doc.add_paragraph(f'- {obj}')

# Módulos y Contenido del Curso
doc.add_heading('Módulos y Contenido del Curso', level=2)

modules = [
    {
        "title": "Módulo 1: Introducción a la Ética y Bioética en la Investigación",
        "topics": [
            "Conceptos básicos de ética y bioética.",
            "Historia y evolución de la bioética.",
            "Principios bioéticos fundamentales."
        ],
        "activities": [
            "Lecturas y análisis de textos clave.",
            "Discusión en foro sobre la relevancia histórica de la bioética."
        ]
    },
    {
        "title": "Módulo 2: Ética en la Investigación con Seres Vivos",
        "topics": [
            "Normativas y estándares éticos.",
            "Consideraciones éticas en la investigación animal.",
            "Alternativas y las 3 Rs (reemplazo, reducción, refinamiento)."
        ],
        "activities": [
            "Caso de estudio sobre un experimento con animales.",
            "Ensayo crítico sobre la aplicación de las 3 Rs."
        ]
    },
    {
        "title": "Módulo 3: Ética en la Investigación sin Seres Vivos",
        "topics": [
            "Normativas aplicables.",
            "Ejemplos de investigaciones y sus consideraciones éticas."
        ],
        "activities": [
            "Debate en foro sobre un caso específico.",
            "Creación de un informe ético para un proyecto de investigación."
        ]
    },
    {
        "title": "Módulo 4: Participación de Seres Humanos en la Investigación",
        "topics": [
            "Derechos y garantías para los participantes.",
            "Consentimiento informado.",
            "Protección de datos y confidencialidad."
        ],
        "activities": [
            "Video explicativo sobre el consentimiento informado.",
            "Desarrollo de un formulario de consentimiento para un estudio hipotético."
        ]
    },
    {
        "title": "Módulo 5: Integridad Científica",
        "topics": [
            "Definición y principios de la integridad científica.",
            "Casos de mala conducta científica.",
            "Políticas y regulaciones internacionales."
        ],
        "activities": [
            "Análisis crítico de un caso de mala conducta científica.",
            "Presentación en grupo sobre cómo fomentar la integridad en un equipo de investigación."
        ]
    }
]

for module in modules:
    doc.add_heading(module["title"], level=3)
    doc.add_paragraph("**Temas:**")
    for topic in module["topics"]:
        doc.add_paragraph(f'- {topic}')
    doc.add_paragraph("**Actividades:**")
    for activity in module["activities"]:
        doc.add_paragraph(f'- {activity}')

# Evaluación del Curso
doc.add_heading('Evaluación del Curso', level=2)

# Actividades de Evaluación
doc.add_heading('Actividades de Evaluación', level=3)
evaluations = [
    "Foro de Discusión: Participación en debates sobre temas éticos.",
    "Video Análisis: Creación de videos cortos aplicando el método deliberativo de Diego Gracia a casos de estudio.",
    "Ensayos y Reportes: Elaboración de ensayos críticos y reportes éticos.",
    "Proyectos Grupales: Presentaciones y trabajos colaborativos."
]
for evaluation in evaluations:
    doc.add_paragraph(f'- {evaluation}')

# Criterios de Evaluación
doc.add_heading('Criterios de Evaluación', level=3)
doc.add_paragraph(
    "- SER: Responsabilidad y entrega puntual de actividades, demostración de interés y pensamiento crítico.\n"
    "- SABER: Comprensión y aplicación de los criterios y estándares de ética, bioética e integridad científica.\n"
    "- HACER: Capacidad de analizar críticamente procesos de investigación y proponer soluciones éticas."
)

# Recursos Adicionales
doc.add_heading('Recursos Adicionales', level=2)
resources = [
    "Bioética e Integridad Científica: http://www.scielo.org.co/scielo.php?pid=S1657-47022018000100006&script=sci_arttext",
    "Política de Ética y Bioética: https://minciencias.gov.co/sites/default/files/upload/noticias/politica-etica.pdf",
    "Bases de Datos y Repositorios Institucionales",
    "Plataforma 'Tu Aula'"
]
for resource in resources:
    doc.add_paragraph(f'- {resource}')

# Cronograma Tentativo
doc.add_heading('Cronograma Tentativo', level=2)
schedule = [
    "Semana 1-2: Introducción y Módulo 1",
    "Semana 3-4: Módulo 2",
    "Semana 5-6: Módulo 3",
    "Semana 7-8: Módulo 4",
    "Semana 9-10: Módulo 5 y Evaluaciones Finales"
]
for week in schedule:
    doc.add_paragraph(f'- {week}')

# Guardar el documento
doc_path = "Modelo_Curso_MOOC_Etica_Bioetica_Investigacion_Integridad_Cientifica.docx"
doc.save(doc_path)

doc_path